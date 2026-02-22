from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def proper_case(name):
    parts = name.strip().split()
    return " ".join(p.capitalize() for p in parts[:2])

def remove_bullets(text):
    return re.sub(r'[•●▪►\-]', '', text)

def generate_docx(text):
    doc = Document()

    lines = text.split("\n")
    name = proper_case(lines[0])
    filename = f"{name}.docx"
    output_path = f"outputs/{filename}"

    # Candidate Name
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # One line space
    doc.add_paragraph("")

    sections = ["Summary", "Technical Skills", 
                "Education, Certification & Training", 
                "Professional Experience"]

    current_section = None
    content = {}

    for line in lines:
        if line.strip() in sections:
            current_section = line.strip()
            content[current_section] = []
        elif current_section:
            content[current_section].append(line)

    for section in sections:
        if section in content:
            doc.add_paragraph("")
            heading = doc.add_paragraph()
            run = heading.add_run(section)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

            for line in content[section]:
                clean = remove_bullets(line).strip()
                if clean:
                    p = doc.add_paragraph(f"• {clean}")
                    p.runs[0].font.name = "Times New Roman"
                    p.runs[0].font.size = Pt(10)

    doc.save(output_path)
    return output_path, filename